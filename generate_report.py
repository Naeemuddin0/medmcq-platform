from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_report():
    doc = Document()

    # --- Header ---
    title = doc.add_heading('COMSATS UNIVERSITY, ISLAMABAD', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Department of Computer Science\nAssignment - 3, Spring 2026\n[CLO4]: Apply DevOps pipeline automation techniques for code deployment')
    run.bold = True

    # --- Student Info ---
    info = doc.add_table(rows=3, cols=2)
    info.cell(0, 0).text = 'Name:'
    info.cell(0, 1).text = 'Naeem ud din'
    info.cell(1, 0).text = 'Registration No:'
    info.cell(1, 1).text = 'SP23-BDS-038'
    info.cell(2, 0).text = 'Instructor:'
    info.cell(2, 1).text = 'Qasim Malik'

    doc.add_paragraph('\n')

    # --- Part I: Automated Test Cases ---
    doc.add_heading('Part-I: Automated Test Cases using Selenium', level=1)
    doc.add_paragraph(
        'In this part, 15 automated test cases were implemented using Python, Selenium, and Pytest. '
        'The tests cover critical UI elements, navigation links, and authentication redirects to ensure the application stability.'
    )

    doc.add_heading('List of Test Cases:', level=2)
    tests = [
        "test_page_title: Verifies the home page title.",
        "test_hero_section_presence: Checks for the main marketing text.",
        "test_nav_about: Verifies navigation to the About page.",
        "test_nav_contact: Verifies navigation to the Contact page.",
        "test_login_page_loads: Checks if the login form is accessible.",
        "test_register_page_loads: Checks if the registration form is accessible.",
        "test_navbar_presence: Ensures the navigation bar is visible.",
        "test_invalid_login_error: Validates form behavior on incorrect credentials.",
        "test_copyright_presence: Checks for the copyright notice in the footer area.",
        "test_logo_link: Verifies the logo redirects to the homepage.",
        "test_practice_redirect: Checks auth protection for the /practice route.",
        "test_dashboard_redirect: Checks auth protection for the /dashboard route.",
        "test_contact_form_inputs: Validates fields in the contact form.",
        "test_meta_description: Verifies SEO meta tags.",
        "test_home_signin_button: Ensures the CTA button is present."
    ]
    for test in tests:
        doc.add_paragraph(test, style='List Bullet')

    doc.add_paragraph('\n[Selenium Test Script - tests/test_medmcq.py]')
    doc.add_paragraph('The script uses robust XPATH selectors and a containerized standalone-chrome driver for stable execution.')

    # --- Part II: Jenkins Pipeline Integration ---
    doc.add_heading('Part-II: Jenkins Pipeline Integration', level=1)
    doc.add_paragraph(
        'A dedicated "Test" stage was integrated into the Jenkinsfile. This stage automates the entire '
        'testing lifecycle: building the test container, running the suite against the dev environment, '
        'and extracting results for reporting.'
    )

    doc.add_heading('Pipeline Stage View:', level=2)
    doc.add_paragraph('The pipeline now includes: Checkout -> Deploy Dev Environment -> Test.')
    doc.add_paragraph('The Test stage is configured to publish JUnit reports, providing interactive visualization of test outcomes in Jenkins.')

    # --- Part III: Email Notifications ---
    doc.add_heading('Part-III: Automated Email Notifications', level=1)
    doc.add_paragraph(
        'Automated email notifications were implemented using the Jenkins Email Extension plugin. '
        'The pipeline is configured to send a detailed build report, including test results as attachments, '
        'to the course instructor (qasimalik@gmail.com) upon completion of the pipeline.'
    )

    # --- Conclusion ---
    doc.add_heading('Conclusion', level=1)
    doc.add_paragraph(
        'The automated testing pipeline successfully achieved an Unstable/Success status, confirming that '
        'the automation logic is sound. 13 out of 15 tests passed in the final verification run (Build #32), '
        'providing immediate feedback on application regressions.'
    )

    doc.save('DevOps_Assignment_3_Report.docx')
    print("Report generated successfully: DevOps_Assignment_3_Report.docx")

if __name__ == '__main__':
    create_report()
