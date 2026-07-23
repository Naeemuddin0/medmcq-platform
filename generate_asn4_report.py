import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor
from docx.oxml.ns import qn

from docx.oxml import OxmlElement

def add_terminal_block(doc, text):
    # Adds a block of text that mimics a terminal
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    
    # Add dark background shading to paragraph
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '282C34')
    pPr.append(shd)
    
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White text

def create_report():
    doc = Document()

    # --- Header ---
    title = doc.add_heading('COMSATS UNIVERSITY, ISLAMABAD', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('Department of Computer Science\nAssignment - 4, Spring 2026\n[CLO5]: Develop cloud native applications using current DevOps tools')
    run.bold = True

    # --- Student Info ---
    info = doc.add_table(rows=3, cols=2)
    info.cell(0, 0).text = 'Name:'
    info.cell(0, 1).text = 'Naeem ud din'
    info.cell(1, 0).text = 'Registration No:'
    info.cell(1, 1).text = 'SP23-BDS-038'
    info.cell(2, 0).text = 'Instructor:'
    info.cell(2, 1).text = 'Qasim Malik'

    doc.add_paragraph('\n')

    # --- Introduction ---
    doc.add_heading('1. Introduction', level=1)
    doc.add_paragraph(
        'This assignment demonstrates the deployment of a cloud-native Next.js web application and a MongoDB database '
        'onto a local Kubernetes cluster using Minikube. The deployment utilizes declarative YAML manifests to '
        'define Pods, Deployments, Services (NodePort), Persistent Volume Claims (PVC), and a Horizontal Pod Autoscaler (HPA).'
    )

    # --- Step 1 ---
    doc.add_heading('2. Kubernetes Manifests', level=1)
    doc.add_paragraph(
        'The application infrastructure was defined using multiple YAML files inside the `k8s/` directory. '
        'This included configuring a PVC to ensure MongoDB data persists across pod restarts, and setting up '
        'NodePort services to expose both the database and the web application.'
    )
    add_terminal_block(doc, "web-deployment.yaml snippet:\napiVersion: apps/v1\nkind: Deployment\nmetadata:\n  name: web-deployment\n...")

    # --- Step 2 ---
    doc.add_heading('3. Minikube Initialization', level=1)
    doc.add_paragraph(
        'Minikube was initialized on the AWS EC2 instance using the Docker driver. Essential addons, such as '
        'the metrics-server (required for HPA) and the dashboard, were enabled.'
    )
    add_terminal_block(doc, 
        "$ minikube start --driver=docker\n"
        "😄  minikube v1.38.1 on Ubuntu 22.04\n"
        "✨  Using the docker driver based on user configuration\n"
        "👍  Starting control plane node minikube in cluster minikube\n"
        "...\n"
        "$ minikube addons enable metrics-server\n"
        "🌟  The 'metrics-server' addon is enabled"
    )

    # --- Step 3 ---
    doc.add_heading('4. Docker Image Build & Deployment', level=1)
    doc.add_paragraph(
        'A lightweight Docker image was built utilizing a pre-compiled Next.js standalone build. The image '
        'was loaded directly into Minikube\'s internal Docker daemon to optimize deployment speed and bypass external registry pulls.'
    )
    add_terminal_block(doc,
        "$ docker build -t medmcq-web:latest -f Dockerfile.fast .\n"
        "$ minikube image load medmcq-web:latest\n"
        "$ kubectl apply -f k8s/\n"
        "deployment.apps/mongo-deployment created\n"
        "persistentvolumeclaim/mongo-pvc created\n"
        "service/mongo-service created\n"
        "deployment.apps/web-deployment created\n"
        "horizontalpodautoscaler.autoscaling/web-hpa created\n"
        "service/web-service created"
    )

    # --- Step 4 ---
    doc.add_heading('5. Verification of Resources', level=1)
    doc.add_paragraph(
        'The status of all deployed resources was verified using kubectl. Both pods successfully transitioned '
        'to the Running state, and the HPA began monitoring CPU utilization.'
    )
    add_terminal_block(doc,
        "$ kubectl get pods,svc,hpa\n"
        "NAME                                READY   STATUS    RESTARTS   AGE\n"
        "pod/mongo-deployment-56d578...      1/1     Running   0          31s\n"
        "pod/web-deployment-766bc75...       1/1     Running   0          31s\n\n"
        "NAME                    TYPE        CLUSTER-IP       PORT(S)\n"
        "service/kubernetes      ClusterIP   10.96.0.1        443/TCP\n"
        "service/mongo-service   NodePort    10.110.220.108   27017:30017/TCP\n"
        "service/web-service     NodePort    10.110.91.12     3000:30080/TCP\n\n"
        "NAME                                          REFERENCE                   TARGETS\n"
        "horizontalpodautoscaler.autoscaling/web-hpa   Deployment/web-deployment   <unknown>/50%"
    )

    # --- Screenshots ---
    doc.add_heading('6. Screenshots', level=1)
    doc.add_paragraph('Below are the screenshots capturing the Minikube dashboard and the deployed Web Application accessed via secure port-forwarding tunnels.')

    BASE_DIR = os.path.dirname(os.path.abspath(__file__))
    images = [
        (os.path.join(BASE_DIR, "asn4 ss", "kubernetes 1.png"), "Minikube Dashboard Overview"),
        (os.path.join(BASE_DIR, "asn4 ss", "kubernetes 2.png"), "Minikube Deployments and Pods"),
        (os.path.join(BASE_DIR, "asn4 ss", "website 3000.png"), "Running Next.js Web Application (Port 3000)")
    ]

    for img_path, caption in images:
        if os.path.exists(img_path):
            doc.add_paragraph(caption, style='Intense Quote')
            doc.add_picture(img_path, width=Inches(6.0))
            doc.add_paragraph('\n')
        else:
            doc.add_paragraph(f"[Screenshot Missing: {img_path}]")

    doc.save('DevOps_Assignment_4_Kubernetes_Report_V3.docx')
    print("Report generated successfully: DevOps_Assignment_4_Kubernetes_Report_V3.docx")

if __name__ == '__main__':
    create_report()
